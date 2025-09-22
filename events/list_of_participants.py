import io
import itertools
from enum import Enum
from typing import Callable, Optional

from django.db.models.functions import Coalesce
from django.utils import timezone
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph

from events.exports.docx_utils import (
    add_page_number,
    insert_hr,
    set_cell,
    set_repeat_table_header,
    set_table_border,
    set_table_caption,
)
from events.models import Registration


class Section(Enum):
    PARTIES = "Parties"
    ASS_PANELS = "Assessment Panels"
    OBSERVERS = "Observers"
    SECRETARIAT = "Ozone Secretariat"


# In order to preserve the layout, use a placeholder if no symbols are provided
# in the event settings.
SYMBOL_PLACEHOLDER = ["UNEP/"]

ZERO = Inches(0)
USABLE_WIDTH = Inches(8)
INDEX_COL_WIDTH = Inches(0.5)
CONTACT_COL_WIDTH = Inches(3.45)

TOP_MARGIN = Inches(0.25)
BOTTOM_MARGIN = Inches(0.25)
LEFT_MARGIN = Inches(0.25)
RIGHT_MARGIN = Inches(0.25)


class ListOfParticipants:
    def __init__(self, event):
        self.doc = Document()
        self.event = event
        self.global_index = 0
        self.sections = {s: [] for s in Section}
        self.get_all_registrations()

    def get_all_registrations(self):
        for r in self.get_query():
            try:
                if r.usable_organization.organization_type.hide_in_lop:
                    continue
            except AttributeError:
                pass

            if r.role.hide_in_lop:
                continue

            if r.is_gov and r.usable_government:
                self.sections[Section.PARTIES].append(r)
            elif r.is_ass_panel:
                self.sections[Section.ASS_PANELS].append(r)
            elif r.is_secretariat:
                self.sections[Section.SECRETARIAT].append(r)
            else:
                self.sections[Section.OBSERVERS].append(r)

    def get_query(self):
        return (
            self.event.registrations.filter(status=Registration.Status.REGISTERED)
            .annotate(actual_sort_order=Coalesce("sort_order", "role__sort_order"))
            .prefetch_related(
                "role",
                "contact",
                "contact__country",
                "contact__organization",
                "contact__organization__organization_type",
                "contact__organization__country",
                "contact__organization__government",
                "organization",
                "organization__organization_type",
                "organization__country",
                "organization__government",
            )
            .order_by("actual_sort_order", "contact__last_name", "contact__first_name")
        )

    def export_docx(self):
        self.init_docx()
        self.get_content()
        return self.save_docx()

    def init_docx(self):
        self.doc.settings.odd_and_even_pages_header_footer = True

        styles = self.doc.styles
        styles["Normal"].font.name = "Arial"

        header_style = styles.add_style("LOP Header", WD_STYLE_TYPE.PARAGRAPH)
        header_style.font.name = "Arial"
        header_style.font.bold = False
        header_style.font.italic = True
        header_style.font.size = Pt(9)
        header_style.paragraph_format.space_after = Pt(0)
        header_style.paragraph_format.space_before = Pt(0)

        header_style = styles.add_style("LOP Header Title", WD_STYLE_TYPE.PARAGRAPH)
        header_style.font.name = "Arial"
        header_style.font.bold = True
        header_style.font.size = Pt(16)
        header_style.paragraph_format.keep_together = True
        header_style.paragraph_format.keep_with_next = True
        header_style.paragraph_format.space_after = Pt(0)
        header_style.paragraph_format.space_before = Pt(0)

        table_style = styles.add_style("LOP Table", WD_STYLE_TYPE.TABLE)
        table_style.font.name = "Arial"
        table_style.font.size = Pt(11)

        space_style = styles.add_style("LOP Table Space", WD_STYLE_TYPE.PARAGRAPH)
        space_style.paragraph_format.space_after = Pt(0)
        space_style.paragraph_format.space_before = Pt(0)
        space_style.font.size = Pt(4)

        group_style = styles.add_style("LOP L1 Group", WD_STYLE_TYPE.PARAGRAPH)
        group_style.font.name = "Arial"
        group_style.font.size = Pt(14)
        group_style.font.bold = True
        group_style.paragraph_format.keep_together = True
        group_style.paragraph_format.keep_with_next = True

        group_style = styles.add_style("LOP L2 Group", WD_STYLE_TYPE.PARAGRAPH)
        group_style.font.name = "Arial"
        group_style.font.bold = True
        group_style.font.underline = True
        group_style.paragraph_format.keep_together = True
        group_style.paragraph_format.keep_with_next = True

        footer_style = styles.add_style("LOP Footer", WD_STYLE_TYPE.PARAGRAPH)
        footer_style.font.name = "Arial"
        footer_style.font.size = Pt(9)
        footer_style.paragraph_format.space_after = Pt(0)
        footer_style.paragraph_format.space_before = Pt(0)

        self.doc.sections[0].different_first_page_header_footer = True
        self.set_footer(self.doc.sections[0].footer)
        self.set_footer(self.doc.sections[0].even_page_footer)

    def set_footer(self, footer):
        footer.paragraphs[0].clear()

        insert_hr(footer.paragraphs[0])
        table = footer.add_table(1, 2, USABLE_WIDTH)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_border(table, size=0, border_type="none")

        p = table.cell(0, 0).paragraphs[0]
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        p.style = "LOP Footer"
        p.add_run(str(timezone.now().date()))

        p = table.cell(0, 1).paragraphs[0]
        p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        p.style = "LOP Footer"
        add_page_number(p.add_run())

    def get_content(self):
        self.cover_page()

        self.grouped_participants(
            section=Section.PARTIES,
            sort=lambda r: (r.usable_government.name,),
            key_l1=None,
            key_l2=lambda r: r.usable_government.name,
        )
        self.grouped_participants(
            section=Section.ASS_PANELS,
            sort=lambda r: (
                r.usable_organization_sort_order,
                (r.usable_organization_name or "Unknown").lower(),
            ),
            key_l1=None,
            key_l2=lambda r: r.usable_organization_name,
        )
        self.grouped_participants(
            section=Section.OBSERVERS,
            sort=lambda r: (
                r.usable_organization_type_sort_order,
                r.usable_organization_type_description or "Unknown",
                r.usable_organization_sort_order,
                (r.usable_organization_name or "Unknown").lower(),
            ),
            key_l1=lambda r: r.usable_organization_type_description or "Unknown",
            key_l2=lambda r: r.usable_organization_name or "Unknown",
        )
        self.grouped_participants(
            section=Section.SECRETARIAT,
            sort=lambda r: (
                r.usable_organization_sort_order,
                r.usable_organization_name.lower(),
            ),
            key_l1=None,
            key_l2=lambda r: r.usable_organization_name,
        )

    def cover_page(self):
        section = self.doc.sections[0]
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

        self.cover_page_header()
        self.cover_page_logos()
        self.cover_page_event()

        p = self.doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        r = p.add_run("\n\n\nList of Participants")
        r.font.size = Pt(32)
        r.font.bold = True

    def cover_page_header(self):
        table = self.doc.add_table(1, 2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style.font.name = "Arial"
        table.style.paragraph_format.space_after = ZERO

        set_table_border(table, size=0, border_type="none")

        set_cell(
            table=table,
            row=0,
            col=0,
            text="UNITED\nNATIONS",
            align=WD_PARAGRAPH_ALIGNMENT.LEFT,
            font_size=Pt(16),
            bold=True,
            space_after=ZERO,
        )
        set_cell(
            table=table,
            row=0,
            col=1,
            text="EP",
            align=WD_PARAGRAPH_ALIGNMENT.RIGHT,
            font_size=Pt(32),
            bold=True,
            space_after=ZERO,
        )

        p = self.doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        is_first = True
        for doc_symbol in self.event.lop_doc_symbols or SYMBOL_PLACEHOLDER:
            if is_first:
                is_first = False
            else:
                p.add_run("\n")

            start, end = doc_symbol.split("/", 1)

            r = p.add_run(start + "/")
            r.font.size = Pt(16)
            r.font.bold = True

            r = p.add_run(end)
            r.font.size = Pt(14)
        insert_hr(p)

    def cover_page_logos(self):
        logo_size = Inches(1)
        col1_width = Inches(1.4)
        col2_width = Inches(5.5)

        table = self.doc.add_table(2, 2)
        table.allow_autofit = False
        table.columns[0].width = col1_width
        table.columns[1].width = col2_width
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_border(table, size=0, border_type="none")

        set_cell(
            table=table,
            row=0,
            col=0,
            image="img/UN-logo.png",
            image_width=logo_size,
            image_height=logo_size,
            align=WD_PARAGRAPH_ALIGNMENT.CENTER,
            v_align=WD_CELL_VERTICAL_ALIGNMENT.CENTER,
            space_after=ZERO,
            cell_width=col1_width,
        )
        set_cell(
            table=table,
            row=0,
            col=1,
            text="\n".join(
                [
                    "Distr.: General",
                    self.event.end_date.strftime("%-d %b %Y"),
                    "",
                    "English only",
                ]
            ),
            align=WD_PARAGRAPH_ALIGNMENT.RIGHT,
            font_size=Pt(12),
            space_after=ZERO,
            cell_width=col2_width,
        )

        set_cell(
            table=table,
            row=1,
            col=0,
            image="img/UNEP-logo.png",
            image_width=logo_size,
            image_height=logo_size,
            align=WD_PARAGRAPH_ALIGNMENT.CENTER,
            v_align=WD_CELL_VERTICAL_ALIGNMENT.CENTER,
            space_after=ZERO,
            cell_width=col1_width,
        )
        set_cell(
            table=table,
            row=1,
            col=1,
            text="\n".join(
                [
                    "United Nations",
                    "Environment",
                    "Programme",
                ]
            ),
            align=WD_PARAGRAPH_ALIGNMENT.LEFT,
            v_align=WD_CELL_VERTICAL_ALIGNMENT.CENTER,
            font_size=Pt(16),
            space_after=ZERO,
            bold=True,
            cell_width=col2_width,
        )

        p = self.doc.add_paragraph()
        insert_hr(p, 32)

    def cover_page_event(self):
        table = self.doc.add_table(2, 2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_border(table, size=0, border_type="none")
        set_cell(
            table=table,
            row=0,
            col=0,
            text=self.event.title,
            font_size=Pt(12),
            space_after=ZERO,
            bold=True,
        )

        set_cell(
            table=table,
            row=1,
            col=0,
            text=", ".join(
                [
                    self.event.venue_city,
                    self.event.dates,
                ]
            ),
            font_size=Pt(12),
            space_after=ZERO,
        )

    def configure_section(self, section, header_text):
        section.different_first_page_header_footer = False
        section.top_margin = TOP_MARGIN
        section.bottom_margin = BOTTOM_MARGIN
        section.left_margin = LEFT_MARGIN
        section.right_margin = RIGHT_MARGIN

        # Set the layout to two columns
        cols = section._sectPr.xpath("./w:cols")[0]
        cols.set(qn("w:num"), "2")
        cols.set(qn("w:space"), "144")

        self.set_section_header(
            section.header, header_text, WD_PARAGRAPH_ALIGNMENT.RIGHT
        )
        self.set_section_header(
            section.even_page_header, header_text, WD_PARAGRAPH_ALIGNMENT.LEFT
        )

    def set_section_header(self, header, header_text, alignment):
        header.is_linked_to_previous = False
        p = header.paragraphs[0]
        p.alignment = alignment
        p.style = "LOP Header"
        p.text = "-".join(self.event.lop_doc_symbols or SYMBOL_PLACEHOLDER)

        insert_hr(p)

        table = header.add_table(1, 1, USABLE_WIDTH)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_cell(
            table=table,
            row=0,
            col=0,
            text=header_text,
            style="LOP Header Title",
            align=WD_PARAGRAPH_ALIGNMENT.CENTER,
            v_align=WD_CELL_VERTICAL_ALIGNMENT.CENTER,
            bg_color="#e0e0e0",
        )
        header.add_paragraph("", style="LOP Table Space")

    def grouped_participants(
        self,
        section: Section,
        sort: Callable[[Registration], tuple],
        key_l1: Optional[Callable[[Registration], str]],
        key_l2: Callable[[Registration], str],
    ):
        if not self.sections[section]:
            return

        self.configure_section(self.doc.add_section(), section.value)

        registrations = sorted(self.sections[section], key=sort)

        if not key_l1:
            key_l1 = lambda x: ""  # noqa: E731

        for l1_group_name, l1_group_items in itertools.groupby(
            registrations, key=key_l1
        ):
            if l1_group_name:
                self.doc.add_paragraph(l1_group_name, style="LOP L1 Group")

            for l2_group_name, l2_group_items in itertools.groupby(
                l1_group_items, key=key_l2
            ):
                items = list(l2_group_items)

                table = self.doc.add_table(len(items) + 1, 2, style="LOP Table")
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.allow_autofit = False
                # Set the column widths here but also for individual cells later.
                # Required since Word doesn't understand column width, unlike other
                # engines like LibreOffice.
                table.columns[0].width = INDEX_COL_WIDTH
                table.columns[1].width = CONTACT_COL_WIDTH
                set_table_border(table, size=0, border_type="none")
                set_table_caption(table, l2_group_name)

                header_row = table.rows[0]
                header_row.cells[0].merge(header_row.cells[1])
                set_repeat_table_header(header_row)

                header_p = table.cell(0, 0).paragraphs[0]
                header_p.style = "LOP L2 Group"
                header_p.add_run(l2_group_name)

                for row_index, item in enumerate(items, start=1):
                    self.global_index += 1

                    # Don't allow rows to be split across pages (or columns
                    row = table.rows[row_index]
                    row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))

                    # Set the global index, tracked across sections
                    cell = row.cells[0]
                    cell.width = INDEX_COL_WIDTH
                    p = cell.paragraphs[0]
                    p.add_run(f"{self.global_index}.").bold = True
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    p.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

                    # Set the contact information
                    cell = row.cells[1]
                    cell.width = CONTACT_COL_WIDTH
                    p = row.cells[1].paragraphs[0]
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    p.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                    self.contact_list_item(p, item, section)

                self.doc.add_paragraph("", style="LOP Table Space")

    def contact_list_item(self, p: Paragraph, reg: Registration, section: Section):
        p.paragraph_format.keep_together = True
        p.add_run(reg.contact.full_name.strip() or "-").bold = True

        organization = reg.organization or reg.contact.organization
        address_obj = (
            organization if reg.contact.is_use_organization_address else reg.contact
        )

        possible_values = [
            reg.designation or reg.contact.designation,
            reg.department or reg.contact.department,
            organization and organization.name
            if section != Section.ASS_PANELS
            else None,
            " ".join(
                filter(
                    lambda x: x,
                    [
                        address_obj.city.strip(),
                        address_obj.state.strip(),
                        address_obj.postal_code.strip(),
                    ],
                )
            ),
            address_obj.country and address_obj.country.name,
        ]

        values = []
        for value in possible_values:
            if value := (value or "").strip():
                values.append(value)

        if values:
            p.add_run("\n" + "\n".join(values))

        emails = [e.strip() for e in reg.contact.emails if e.strip()]
        if emails:
            p.add_run("\nEmail: " + ", ".join(emails))

    def save_docx(self):
        doc_file = io.BytesIO()
        doc_file.name = f"{self.event.code}-LoP.docx"
        self.doc.save(doc_file)
        doc_file.seek(0)
        return doc_file
