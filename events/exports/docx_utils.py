from django.contrib.staticfiles import finders
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, ns
from docx.oxml.ns import qn
from docx.shared import Length
from docx.table import Table
from docx.text.paragraph import Paragraph


def set_repeat_table_header(row):
    """Set repeat table row on every new page"""
    tr = row._tr
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr.get_or_add_trPr().append(tbl_header)
    return row


def set_cell(
    table: Table,
    row: int,
    col: int,
    text: str = "",
    image: str = None,
    image_width: Length = None,
    image_height: Length = None,
    style: str = None,
    align: WD_PARAGRAPH_ALIGNMENT = None,
    v_align: WD_CELL_VERTICAL_ALIGNMENT = None,
    bg_color: str = None,
    font_size: Length = None,
    bold: bool = None,  # noqa: FBT001
    space_after: Length = None,
    cell_width: Length = None,
):
    cell = table.cell(row, col)
    cell.text = ""

    if cell_width:
        cell.width = cell_width

    p = cell.paragraphs[0]
    run = p.add_run(text)
    if image:
        run.add_picture(finders.find(image), width=image_width, height=image_height)

    p.style = style
    p.alignment = align

    if font_size:
        run.font.size = font_size
    if bold is not None:
        run.font.bold = bold
    if space_after is not None:
        p.paragraph_format.space_after = space_after

    if bg_color:
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), bg_color)
        cell._tc.get_or_add_tcPr().append(shd)
    cell.vertical_alignment = v_align
    return cell


def set_table_border(
    table: Table, color="000000", size=4, space=0, border_type="single"
):
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    # Remove existing borders if present
    borders = tbl_pr.xpath("w:tblBorders")
    if borders:
        for b in borders:
            tbl_pr.remove(b)

    tbl_borders = OxmlElement("w:tblBorders")

    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), border_type)
        border.set(qn("w:sz"), str(size))
        border.set(qn("w:space"), str(space))
        border.set(qn("w:color"), color)
        tbl_borders.append(border)

    tbl_pr.append(tbl_borders)
    return table


def set_table_caption(table: Table, name: str):
    tbl_pr = table._tbl.tblPr

    tbl_name = OxmlElement("w:tblCaption")
    tbl_name.set(qn("w:val"), name)
    tbl_pr.append(tbl_name)
    return table


def insert_hr(paragraph: Paragraph, size=6):
    p = paragraph._p  #
    p_pr = p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    p_pr.insert_element_before(
        p_bdr,
        "w:shd",
        "w:tabs",
        "w:suppressAutoHyphens",
        "w:kinsoku",
        "w:wordWrap",
        "w:overflowPunct",
        "w:topLinePunct",
        "w:autoSpaceDE",
        "w:autoSpaceDN",
        "w:bidi",
        "w:adjustRightInd",
        "w:snapToGrid",
        "w:spacing",
        "w:ind",
        "w:contextualSpacing",
        "w:mirrorIndents",
        "w:suppressOverlap",
        "w:jc",
        "w:textDirection",
        "w:textAlignment",
        "w:textboxTightWrap",
        "w:outlineLvl",
        "w:divId",
        "w:cnfStyle",
        "w:rPr",
        "w:sectPr",
        "w:pPrChange",
    )
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    p_bdr.append(bottom)


def create_element(name):
    return OxmlElement(name)


def create_attribute(element, name, value):
    element.set(ns.qn(name), value)


def add_page_number(run):
    fld_char1 = create_element("w:fldChar")
    create_attribute(fld_char1, "w:fldCharType", "begin")

    instr_text = create_element("w:instrText")
    create_attribute(instr_text, "xml:space", "preserve")
    instr_text.text = "PAGE"

    fld_char2 = create_element("w:fldChar")
    create_attribute(fld_char2, "w:fldCharType", "end")

    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    return run
