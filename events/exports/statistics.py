import io
from collections import Counter
from decimal import Decimal
from typing import Callable, Collection, Iterable

from django.utils import timezone
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Inches, Pt, RGBColor

from core.models import BaseContact, Country, OrganizationType, Region
from events.exports.docx_utils import set_cell, set_table_border, set_table_caption
from events.models import Event, Registration


class StatisticsBase:
    STATUS = []
    SUFFIX = ""
    TITLE = ""

    def __init__(self, event: Event):
        self.table_count = 0
        self.doc = Document()
        self.event = event
        self.regions = list(
            Region.objects.all()
            .prefetch_related("countries", "countries__subregion")
            .order_by("sort_order", "name")
        )
        self.registrations = []
        for r in self.get_query():
            try:
                if r.usable_organization.organization_type.hide_in_statistics:
                    continue
            except AttributeError:
                pass
            self.registrations.append(r)

        self.gov_registrations = [
            registration for registration in self.registrations if registration.is_gov
        ]
        self.gov_parties = {
            registration.usable_government for registration in self.gov_registrations
        }
        self.org_type_names = [
            org_type.statistics_display_name
            for org_type in OrganizationType.objects.filter(
                hide_in_statistics=False
            ).order_by("sort_order", "statistics_title", "title")
        ]

    def get_query(self):
        query = (
            self.event.registrations.all()
            .prefetch_related(
                "organization",
                "organization__organization_type",
                "organization__government",
                "organization__government__region",
                "organization__government__subregion",
                "contact",
                "contact__organization",
                "contact__organization__organization_type",
                "contact__organization__government",
                "contact__organization__government__region",
                "contact__organization__government__subregion",
            )
            .order_by("contact__last_name", "contact__first_name")
        )
        if self.STATUS:
            query = query.filter(status__in=self.STATUS)

        return query

    def export_docx(self):
        self.init_docx()
        self.get_content()
        return self.save_docx()

    def init_docx(self):
        styles = self.doc.styles

        styles["Normal"].font.name = "Times New Roman"

        table_style = styles.add_style("Stat Table", WD_STYLE_TYPE.TABLE)
        table_style.font.size = Pt(11)
        table_style.font.name = "Times New Roman"
        table_style.paragraph_format.keep_together = True
        table_style.paragraph_format.keep_with_next = True
        table_style.paragraph_format.space_after = Pt(5)
        table_style.paragraph_format.space_before = Pt(5)
        table_style.paragraph_format.left_indent = Inches(0.1)
        table_style.paragraph_format.right_indent = Inches(0.1)

        th_style = styles.add_style("Table Header", WD_STYLE_TYPE.PARAGRAPH)
        th_style.font.bold = True
        th_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        td_style = styles.add_style("Table Data", WD_STYLE_TYPE.PARAGRAPH)
        td_style.font.bold = False

        tf_style = styles.add_style("Table Footer", WD_STYLE_TYPE.PARAGRAPH)
        tf_style.font.bold = True

        for section in self.doc.sections:
            header = section.header.paragraphs[0]
            header.text = str(self.event)
            header.style.font.size = Pt(9)
            header.style.font.color.rgb = RGBColor(45, 45, 45)
            header.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

            footer = section.footer.paragraphs[0]
            footer.text = "Date: " + timezone.now().strftime(" %Y-%m-%d %H:%M %Z")
            footer.style.font.size = Pt(9)
            footer.style.font.color.rgb = RGBColor(45, 45, 45)

            section.top_margin = Cm(0)
            section.bottom_margin = Cm(0)
            section.left_margin = Cm(1)
            section.right_margin = Cm(1)

        self.doc.add_paragraph()

        p = self.doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        p.add_run(timezone.now().date().strftime("%-d %B %Y"))

        if self.TITLE:
            p = self.doc.add_paragraph()
            r = p.add_run(self.TITLE)
            r.font.bold = True

    def get_content(self):
        pass

    def save_docx(self):
        doc_file = io.BytesIO()
        doc_file.name = f"{self.event.code}{self.SUFFIX}.docx"
        self.doc.save(doc_file)
        doc_file.seek(0)
        return doc_file

    def _max_cols(self, obj: Collection[Collection]):
        return max(map(len, obj), default=0)

    def table(
        self,
        name: str,
        head: Collection[Collection],
        body: Collection[Collection],
        footer: Collection[Collection],
    ):
        self.table_count += 1
        name = f"Table {self.table_count}: {name}"

        self.doc.add_paragraph()

        # Add one extra row for the heading with the name
        rows = len(head) + len(body) + len(footer) + 1
        cols = max(map(self._max_cols, (head, body, footer)), default=0)
        if not cols:
            raise RuntimeError("No columns found in head, body or footer")

        table = self.doc.add_table(rows, cols, style="Stat Table")
        set_table_border(table)
        set_table_caption(table, name)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        self.col_span(table, 0, 0, cols)
        self.th(table, 0, 0, name)

        row_index = 1
        for part, func in (head, self.th), (body, self.td), (footer, self.tf):
            for row in part:
                for col_index, data in enumerate(row):
                    if data is not None:
                        func(table, row_index, col_index, str(data))
                row_index += 1

        return table

    def th(self, table, row, col, text):
        return set_cell(
            table=table,
            row=row,
            col=col,
            text=text,
            style="Table Header",
            align=WD_PARAGRAPH_ALIGNMENT.CENTER if row != 0 else None,
            v_align=WD_CELL_VERTICAL_ALIGNMENT.CENTER,
            bg_color="#a5c9eb",
        )

    def td(self, table, row, col, text):
        return set_cell(
            table=table,
            row=row,
            col=col,
            text=str(text),
            style="Table Data",
            align=WD_PARAGRAPH_ALIGNMENT.RIGHT if col != 0 else None,
        )

    def tf(self, table, row, col, text):
        return set_cell(
            table=table,
            row=row,
            col=col,
            text=str(text),
            style="Table Footer",
            align=WD_PARAGRAPH_ALIGNMENT.RIGHT if col != 0 else None,
        )

    def row_span(self, table, row, col, span):
        cell = table.cell(row, col)
        for i in range(row + 1, row + span):
            cell.merge(table.cell(i, col))

    def col_span(self, table, row, col, span):
        cell = table.cell(row, col)
        for i in range(col + 1, col + span):
            cell.merge(table.cell(row, i))

    def _format_percent(self, value, total):
        if total == 0:
            return "-"
        value = Decimal(value) / Decimal(total) * 100
        return str(value.quantize(Decimal("0.01"))) + "%"

    def group_by_org_type(self, registrations: Iterable[Registration]):
        result = {t: [] for t in self.org_type_names}
        for registration in registrations:
            try:
                org = registration.usable_organization
                name = org.organization_type.statistics_display_name
            except AttributeError:
                name = "Unknown"
            if name not in result:
                result[name] = []
            result[name].append(registration)
        return result

    def count_by_gender(self, registrations: Iterable[Registration]):
        result = dict.fromkeys(BaseContact.GenderChoices, 0)
        result["N/A"] = 0
        for registration in registrations:
            result[registration.contact.gender or "N/A"] += 1
        return result


class PreMeetingStatistics(StatisticsBase):
    STATUS = [Registration.Status.ACCREDITED]
    SUFFIX = "-pre-meeting-statistics"
    TITLE = "Meeting Participants statistics: Registration and A5 funding"

    def get_content(self):
        self.table_pax_by_category()
        self.table_parties_by_region()

        for region in self.regions:
            self.table_parties_by_subregion(region)

    def table_pax_by_category(self):
        counts = self.group_by_org_type(self.registrations)

        self.table(
            "Registered participants",
            [
                ("Category", "Registered"),
            ],
            [(category, len(items)) for category, items in counts.items()],
            [
                ("Total", len(self.registrations)),
            ],
        )

    def table_hl(self):
        body = [
            (
                r.usable_government and r.usable_government.name,
                r.usable_organization_name,
                r.contact.full_name,
            )
            for r in self.registrations
            if r.contact.title in BaseContact.HL_TITLES
        ]
        body.sort(key=lambda x: x[0] or "")

        table = self.table(
            "HL Participants Accredited",
            [
                ("Party", "Organization", "Name"),
            ],
            body,
            [[f"Total: {len(body)}"]],
        )
        self.col_span(table, len(body) + 2, 0, 3)

    def table_parties_by_region(self):
        total = len(self.gov_parties)
        counts = dict.fromkeys(self.regions, 0)
        for party in self.gov_parties:
            counts[party.region] += 1

        self.table(
            "Registered Parties",
            [
                ("Category", "Registered", "% of total", "% in category"),
            ],
            [
                (
                    f"{region} Parties",
                    count,
                    self._format_percent(count, total),
                    self._format_percent(count, region.countries.count()),
                )
                for region, count in counts.items()
            ],
            [
                ("Total", total, "", ""),
            ],
        )

    def table_parties_by_subregion(self, region):
        all_parties = Counter(
            (
                country.subregion
                for country in region.countries.all().order_by(
                    "subregion__sort_order", "subregion__name"
                )
            )
        )
        accredited_parties = Counter(
            (
                country.subregion
                for country in self.gov_parties
                if country.region == region
            )
        )
        accredited_registrations = Counter(
            (
                registration.usable_government.subregion
                for registration in self.gov_registrations
                if registration.usable_government.region == region
            )
        )
        table = self.table(
            f"{region} Parties",
            [
                (
                    "Region",
                    "Parties per region",
                    "Parties registered",
                    "Parties to register",
                    "Participants registered",
                    "Funding Requested",
                    None,
                    "Reserved Slots",
                    "Remarks",
                ),
                (None, "(a)", "(b)", "(a - b)", None, "Received", "Approved"),
            ],
            [
                (
                    subregion.name,
                    count,
                    accredited_parties[subregion],
                    count - accredited_parties[subregion],
                    accredited_registrations[subregion],
                )
                for subregion, count in all_parties.items()
            ],
            [
                (
                    "Total",
                    sum(all_parties.values()),
                    sum(accredited_parties.values()),
                    sum(all_parties.values()) - sum(accredited_parties.values()),
                    sum(accredited_registrations.values()),
                )
            ],
        )

        self.row_span(table, 1, 0, 2)
        self.row_span(table, 1, 4, 2)
        self.col_span(table, 1, 5, 2)
        self.row_span(table, 1, 7, 2)
        self.row_span(table, 1, 8, 2)


class PostMeetingStatistics(StatisticsBase):
    STATUS = [Registration.Status.ACCREDITED, Registration.Status.REGISTERED]
    SUFFIX = "-post-meeting-statistics"

    def __init__(self, event: Event):
        super().__init__(event)
        self.registrations_reg = [
            r for r in self.registrations if r.status == r.Status.REGISTERED
        ]
        self.gov_registrations_reg = [
            r for r in self.gov_registrations if r.status == r.Status.REGISTERED
        ]

    def init_docx(self):
        """Overridden to remove document header for post-meeting statistics."""
        super().init_docx()
        for section in self.doc.sections:
            header = section.header.paragraphs[0]
            header.clear()

    def get_content(self):
        self.table_participants_by_gender()
        self.table_hl()

        regions = [r.name for r in self.regions]
        self.table_gov_participants_by_discriminator(
            "Participants Registered",
            "Parties",
            lambda r: r.usable_government.region.name,
            regions,
            self.gov_registrations_reg,
        )
        self.table_parties_by_discriminator(
            "Parties Registered",
            "Parties",
            lambda c: c.region.name,
            regions,
            self.gov_registrations_reg,
        )

        for region in self.regions:
            regs = [
                r
                for r in self.gov_registrations_reg
                if r.usable_government and r.usable_government.region == region
            ]
            subregions = list(
                Country.objects.filter(region=region)
                .values_list("subregion__name", flat=True)
                .order_by("subregion__sort_order", "subregion__name")
            )

            self.table_gov_participants_by_discriminator(
                f"{region} Participants Registered",
                "Region",
                lambda r: r.usable_government.subregion.name,
                subregions,
                regs,
            )
            self.table_parties_by_discriminator(
                f"{region} Parties Registered",
                "Region",
                lambda c: c.subregion.name,
                subregions,
                regs,
            )

    def table_participants_by_gender(self):
        both_by_org_type = self.group_by_org_type(self.registrations)
        reg_by_org_type = self.group_by_org_type(self.registrations_reg)

        body = []
        for org_type, all_regs in both_by_org_type.items():
            regs = reg_by_org_type.get(org_type, [])

            body.append(
                (
                    org_type,
                    len(all_regs),
                    len(regs),
                    *self.count_by_gender(regs).values(),
                )
            )

        table = self.table(
            "Participants by gender",
            [
                ("Category", "Participants Accredited", "Participants Registered"),
                (None, None, "Total", *BaseContact.GenderChoices, "N/A"),
            ],
            body,
            [
                (
                    "Total",
                    len(self.registrations),
                    len(self.registrations_reg),
                    *self.count_by_gender(self.registrations_reg).values(),
                ),
            ],
        )

        self.row_span(table, 1, 0, 2)
        self.row_span(table, 1, 1, 2)
        self.col_span(table, 1, 2, len(BaseContact.GenderChoices) + 2)

    def table_hl(self):
        body = [
            (
                r.usable_government and r.usable_government.name,
                r.usable_organization_name,
                r.contact.full_name,
            )
            for r in self.registrations_reg
            if r.contact.title in BaseContact.HL_TITLES
        ]
        body.sort(key=lambda x: x[0] or "")

        table = self.table(
            "HL Participants Registered",
            [
                ("Party", "Organization", "Name"),
            ],
            body,
            [[f"Total: {len(body)}"]],
        )

        self.col_span(table, len(body) + 2, 0, 3)

    def table_gov_participants_by_discriminator(
        self,
        name: str,
        header: str,
        key_func: Callable[[Registration], str],
        all_values: Collection[str],
        registrations: Collection[Registration],
    ):
        grouped = {v: [] for v in all_values}
        for r in registrations:
            grouped[key_func(r)].append(r)

        body = []
        for value, regs in grouped.items():
            regs = list(regs)
            body.append(
                (
                    value,
                    len(regs),
                    *self.count_by_gender(regs).values(),
                )
            )

        table = self.table(
            name,
            [
                (header, "Participants Registered"),
                (None, "Total", *BaseContact.GenderChoices, "N/A"),
            ],
            body,
            [
                (
                    "Total",
                    len(registrations),
                    *self.count_by_gender(registrations).values(),
                )
            ],
        )
        self.row_span(table, 1, 0, 2)
        self.col_span(table, 1, 1, len(BaseContact.GenderChoices) + 2)

    def table_parties_by_discriminator(
        self,
        name: str,
        header: str,
        key_func: Callable[[Country], str],
        all_values: Collection[str],
        registrations: Collection[Registration],
    ):
        parties = {r.usable_government for r in registrations if r.usable_government}
        counts = dict.fromkeys(all_values, 0)
        for party in parties:
            counts[key_func(party)] += 1

        self.table(
            name,
            [(header, "Parties")],
            [(name, value) for name, value in counts.items()],
            [
                ("Total", len(parties)),
            ],
        )
