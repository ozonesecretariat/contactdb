import docx.document
from docx.blkcntnr import BlockItemContainer
from docx.oxml.ns import qn


def get_table_data(doc: docx.document.Document) -> dict[str, list[str]]:
    tables_data = {}
    for table in doc.tables:
        if (table_pr := table._tbl.tblPr) is None:
            continue
        if (tbl_caption := table_pr.find(qn("w:tblCaption"))) is None:
            continue

        name = tbl_caption.get(qn("w:val"))

        rows = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            rows.append(row_data)
        tables_data[name] = rows
    return tables_data


def extract_text_from_docx_headers(doc: docx.document.Document) -> str:
    result = []
    for section in doc.sections:
        for h in (section.header, section.even_page_header, section.first_page_header):
            result.append(extract_text_from_docx(h))
    return "\n".join(result)


def extract_text_from_docx_footers(doc: docx.document.Document) -> str:
    result = []
    for section in doc.sections:
        for h in (section.footer, section.even_page_footer, section.first_page_footer):
            result.append(extract_text_from_docx(h))
    return "\n".join(result)


def extract_text_from_docx(doc: BlockItemContainer | docx.document.Document) -> str:
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = " ".join(p.text for p in cell.paragraphs)
                if cell_text.strip():
                    full_text.append(cell_text)

    return "\n".join(full_text)
