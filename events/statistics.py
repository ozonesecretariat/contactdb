import io
from collections import Counter
from decimal import Decimal

from django.utils import timezone
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.table import Table

from core.models import OrganizationType, Region
from events.models import Event, Registration


def set_cell(
    table: Table,
    row: int,
    col: int,
    text: str,
    style: str = None,
    align: WD_PARAGRAPH_ALIGNMENT = None,
    v_align: WD_CELL_VERTICAL_ALIGNMENT = None,
    bg_color: str = None,
):
    cell = table.cell(row, col)
    cell.text = ""

    p = cell.paragraphs[0]
    p.add_run(text)

    p.style = style
    p.alignment = align
    if bg_color:
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), bg_color)
        cell._tc.get_or_add_tcPr().append(shd)
    cell.vertical_alignment = v_align
    return cell


def set_table_border(
    table: Table, color="000000", size=4, space=0, border_type="single"
):
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    # Remove existing borders if present
    borders = tbl_pr.xpath("w:tblBorders")
    if borders:
        for b in borders:
            tbl_pr.remove(b)

    tbl_borders = OxmlElement("w:tblBorders")

    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), border_type)  # border style
        border.set(qn("w:sz"), str(size))  # width
        border.set(qn("w:space"), str(space))  # spacing
        border.set(qn("w:color"), color)  # color
        tbl_borders.append(border)

    tbl_pr.append(tbl_borders)
    return table


class PreMeetingStatistics:
    def __init__(self, event: Event):
        self.doc = Document()
        self.event = event
        self.regions = list(
            Region.objects.all()
            .prefetch_related("countries", "countries__subregion")
            .order_by("-name")
        )
        self.accredited_registrations = list(
            event.registrations.filter(
                status=Registration.Status.ACCREDITED
            ).prefetch_related(
                "organization",
                "organization__organization_type",
                "organization__government",
                "organization__government__region",
                "organization__government__subregion",
                "contact",
                "contact__organization",
                "contact__organization__organization_type",
                "contact__organization__government",
                "contact__organization__government__region",
                "contact__organization__government__subregion",
            )
        )
        self.accredited_gov_registrations = [
            registration
            for registration in self.accredited_registrations
            if registration.is_gov
        ]
        self.accredited_gov_parties = {
            registration.usable_government
            for registration in self.accredited_gov_registrations
        }

    def export_docx(self):
        self.init_docx()
        self.get_content()
        return self.save_docx()

    def init_docx(self):
        styles = self.doc.styles
        table_style = styles.add_style("Stat Table", WD_STYLE_TYPE.TABLE)
        table_style.font.size = Pt(11)
        table_style.font.name = "Arial"
        table_style.paragraph_format.keep_together = True
        table_style.paragraph_format.keep_with_next = True
        table_style.paragraph_format.space_after = Pt(0)

        th_style = styles.add_style("Table Header", WD_STYLE_TYPE.PARAGRAPH)
        th_style.font.bold = True
        th_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        td_style = styles.add_style("Table Data", WD_STYLE_TYPE.PARAGRAPH)
        td_style.font.bold = False

        tf_style = styles.add_style("Table Footer", WD_STYLE_TYPE.PARAGRAPH)
        tf_style.font.bold = True

        for section in self.doc.sections:
            header = section.header.paragraphs[0]
            header.text = str(self.event)
            header.style.font.size = Pt(9)
            header.style.font.color.rgb = RGBColor(45, 45, 45)

            footer = section.footer.paragraphs[0]
            footer.text = "Date: " + timezone.now().strftime(" %Y-%m-%d %H:%M %Z")
            footer.style.font.size = Pt(9)
            footer.style.font.color.rgb = RGBColor(45, 45, 45)

            section.top_margin = Cm(0)
            section.bottom_margin = Cm(0)
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(1.5)

    def table(self, rows, cols):
        table = self.doc.add_table(rows, cols, style="Stat Table")
        return set_table_border(table)

    def th(self, table, row, col, text):
        return set_cell(
            table=table,
            row=row,
            col=col,
            text=text,
            style="Table Header",
            v_align=WD_CELL_VERTICAL_ALIGNMENT.CENTER,
            bg_color="#a5c9eb",
        )

    def td(self, table, row, col, text):
        return set_cell(
            table=table,
            row=row,
            col=col,
            text=str(text),
            style="Table Data",
            align=WD_PARAGRAPH_ALIGNMENT.RIGHT if col != 0 else None,
        )

    def tf(self, table, row, col, text):
        return set_cell(
            table=table,
            row=row,
            col=col,
            text=str(text),
            style="Table Footer",
            align=WD_PARAGRAPH_ALIGNMENT.RIGHT if col != 0 else None,
        )

    def get_content(self):
        self.doc.add_paragraph()

        self.table_pax_by_category()
        self.doc.add_paragraph()

        self.table_parties_by_region()
        self.doc.add_paragraph()

        for region in self.regions:
            self.table_parties_by_subregion(
                region, with_funding=region.code.lower() == "a5"
            )
            self.doc.add_paragraph()

    def _format_percent(self, value, total):
        if total == 0:
            return "-"
        value = Decimal(value) / Decimal(total) * 100
        return str(value.quantize(Decimal("0.01"))) + "%"

    def table_pax_by_category(self):
        counts = dict.fromkeys(
            (
                org_type.statistics_display_name
                for org_type in OrganizationType.objects.all()
            ),
            0,
        )
        for registration in self.accredited_registrations:
            try:
                org = registration.usable_organization
                name = org.organization_type.statistics_display_name
            except AttributeError:
                name = "Unknown"
            if name not in counts:
                counts[name] = 0
            counts[name] += 1

        table = self.table(len(counts) + 2, 2)
        self.th(table, 0, 0, "Category")
        self.th(table, 0, 1, "Pax registered")

        for row, (category, count) in enumerate(
            sorted(counts.items(), key=lambda t: t[0]), start=1
        ):
            self.td(table, row, 0, category)
            self.td(table, row, 1, count)

        self.tf(table, len(counts) + 1, 0, "Total")
        self.tf(table, len(counts) + 1, 1, str(len(self.accredited_registrations)))

    def table_parties_by_region(self):
        total = len(self.accredited_gov_parties)
        counts = dict.fromkeys(self.regions, 0)
        for party in self.accredited_gov_parties:
            counts[party.region] += 1

        table = self.table(len(counts) + 2, 4)

        table.cell(0, 0).merge(table.cell(0, 1))
        self.th(table, 0, 0, "Parties registered")
        self.th(table, 0, 2, "% of total")
        self.th(table, 0, 3, "% in category")

        for row, (region, count) in enumerate(counts.items(), start=1):
            self.td(table, row, 0, f"{region} Parties")
            self.td(table, row, 1, count)
            self.td(table, row, 2, self._format_percent(count, total))
            self.td(
                table, row, 3, self._format_percent(count, region.countries.count())
            )

        self.tf(table, len(counts) + 1, 0, "Total")
        self.tf(table, len(counts) + 1, 1, total)

    def table_parties_by_subregion(self, region, with_funding=False):
        all_parties = Counter(
            (
                country.subregion
                for country in region.countries.all().order_by("subregion__name")
            )
        )
        accredited_parties = Counter(
            (
                country.subregion
                for country in self.accredited_gov_parties
                if country.region == region
            )
        )
        accredited_registrations = Counter(
            (
                registration.usable_government.subregion
                for registration in self.accredited_gov_registrations
                if registration.usable_government.region == region
            )
        )

        table = self.table(len(all_parties) + 3, 8 if with_funding else 5)

        table.cell(0, 0).merge(table.cell(1, 0))
        self.th(table, 0, 0, f"Region - {region}")
        self.th(table, 0, 1, "Parties/region")
        self.th(table, 1, 1, "(a)")
        self.th(table, 0, 2, "Parties registered")
        self.th(table, 1, 2, "(b)")
        self.th(table, 0, 3, "Parties to register")
        self.th(table, 1, 3, "(a - b)")
        table.cell(0, 4).merge(table.cell(1, 4))
        self.th(table, 0, 4, "Pax registered")
        if with_funding:
            table.cell(0, 5).merge(table.cell(0, 6))
            self.th(table, 0, 5, "Funding Requested")
            self.th(table, 1, 5, "Received")
            self.th(table, 1, 6, "Approved")
            table.cell(0, 7).merge(table.cell(1, 7))
            self.th(table, 0, 7, "Reserved Slots")

        for row, (subregion, count) in enumerate(all_parties.items(), start=2):
            self.td(table, row, 0, subregion.name)
            self.td(table, row, 1, count)
            self.td(table, row, 2, accredited_parties[subregion])
            self.td(table, row, 3, count - accredited_parties[subregion])
            self.td(table, row, 4, accredited_registrations[subregion])

        row = len(all_parties) + 2
        self.tf(table, row, 0, "Total")
        self.tf(table, row, 1, sum(all_parties.values()))
        self.tf(table, row, 2, sum(accredited_parties.values()))
        self.tf(
            table,
            row,
            3,
            sum(all_parties.values()) - sum(accredited_parties.values()),
        )
        self.tf(table, row, 4, sum(accredited_registrations.values()))

    def save_docx(self):
        doc_file = io.BytesIO()
        doc_file.name = f"{self.event.code}-statistics.docx"
        self.doc.save(doc_file)
        doc_file.seek(0)
        return doc_file
