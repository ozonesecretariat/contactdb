import io
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, RGBColor
from import_export.formats.base_formats import Format


class DOCX(Format):
    def get_title(self):
        return "docx"

    def can_export(self):
        return True

    def get_content_type(self):
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    def get_extension(self):
        return "docx"

    def is_binary(self):
        return False

    def export_data(self, dataset, **kwargs):
        doc = Document()

        section_header = doc.sections[0].header.paragraphs[0]
        section_header.text = "Ozone Contact DB\t"
        section_header.style.font.size = Pt(9)
        section_header.style.font.color.rgb = RGBColor(212, 212, 212)

        styles = doc.styles
        p_style = styles.add_style("FieldStyle", WD_STYLE_TYPE.PARAGRAPH)
        p_style.font.size = Pt(10)
        p_style.font.name = "Arial"
        p_style.font.italic = True
        p_style.paragraph_format.space_after = Pt(1.2)

        r_style = styles.add_style("RunStyle", WD_STYLE_TYPE.CHARACTER)
        r_style.font.size = Pt(10)
        r_style.font.name = "Arial"
        r_style.font.bold = True
        r_style.font.italic = False

        headers = [name.title().replace("_", " ") for name in dataset.headers]

        for row in dataset:
            for header, value in zip(headers, row):
                paragraph = doc.add_paragraph(header + ": ", style="FieldStyle")
                paragraph.add_run(value, style="RunStyle")
            doc.add_page_break()

        doc_file = io.BytesIO()
        doc.save(doc_file)
        return doc_file.getvalue()
