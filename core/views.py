import io

import django.utils.datastructures
import openpyxl as opx
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from django.shortcuts import redirect
from django.core.mail import send_mail
from django.contrib import messages
from django.contrib.auth.mixins import LoginRequiredMixin, PermissionRequiredMixin
from django.http import Http404, HttpResponse, HttpResponseRedirect
from django.urls import reverse
from django.views.generic import (
    View,
    TemplateView,
    DetailView,
    DeleteView,
    UpdateView,
    ListView,
)
from django.views.generic.edit import FormMixin, CreateView, FormView

from core.forms import (
    RecordUpdateForm,
    GroupUpdateForm,
    AddGroupMemberForm,
    AddMultipleGroupMembersForm,
    SendEmailForm,
    KronosEventsImportForm,
    KronosParticipantsImportForm,
    ResolveAllConflictsForm,
    ResolveConflictForm,
    MergeContactsFirstStepForm,
    MergeContactsSecondStepForm,
)

from django_tables2 import SingleTableMixin, SingleTableView
from django_filters.views import FilterView

from core.models import (
    Record,
    RegistrationStatus,
    Group,
    Emails,
    LoadKronosEventsTask,
    LoadKronosParticipantsTask,
    KronosEvent,
    ResolveAllConflictsTask,
    TemporaryContact,
    SendMailTask,
    EmailTemplate,
    EmailFile,
)
from core.tables import (
    RecordTable,
    GroupTable,
    GroupMemberTable,
    LoadKronosEventsTable,
    LoadKronosParticipantsTable,
)
from core.filters import (
    RecordFilter,
    RegistrationStatusFilter,
    GroupFilter,
    GroupMembersFilter,
    SearchContactFilter,
    EmailFilter,
)
from core.utils import ConflictResolutionMethods, update_object


class HomepageView(LoginRequiredMixin, TemplateView):
    template_name = "core/home_page.html"


class RecordDetailView(LoginRequiredMixin, DetailView):
    model = Record


class RecordUpdateView(LoginRequiredMixin, PermissionRequiredMixin, UpdateView):
    model = Record
    form_class = RecordUpdateForm

    def has_permission(self):
        return self.request.user.can_edit

    def get_success_url(self):
        return reverse("contact-detail", kwargs={"pk": self.object.pk})

    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        if self.object.main_contact is not None:
            kwargs["main_contact_choices"] = [
                (None, "---------"),
                (self.object.main_contact.id, self.object.main_contact),
            ]
        else:
            kwargs["main_contact_choices"] = [(None, "---------")]
        return kwargs


class RecordCreateView(LoginRequiredMixin, PermissionRequiredMixin, CreateView):
    model = Record
    form_class = RecordUpdateForm

    def has_permission(self):
        return self.request.user.can_edit

    def get_success_url(self):
        return reverse("contact-detail", kwargs={"pk": self.object.pk})

    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        kwargs["main_contact_choices"] = [(None, "---------")]
        return kwargs


class RecordDeleteView(LoginRequiredMixin, PermissionRequiredMixin, DeleteView):
    model = Record

    def get_success_url(self):
        return reverse("contact-list")

    def has_permission(self):
        return self.request.user.can_edit


class ContactListView(
    LoginRequiredMixin, PermissionRequiredMixin, SingleTableMixin, FilterView
):
    table_class = RecordTable
    queryset = Record.objects.all()
    filterset_class = RecordFilter
    paginate_by = 30

    def has_permission(self):
        return self.request.user.can_view

    def get_template_names(self):
        if self.request.htmx:
            template_name = "records_list_page_partial.html"
        else:
            template_name = "records_list_page.html"

        return template_name

    def get(self, request, *args, **kwargs):
        return super().get(request, *args, **kwargs)


class RegistrationStatusListView(LoginRequiredMixin, FilterView, ListView):
    model = RegistrationStatus
    context_object_name = "statuses"
    filterset_class = RegistrationStatusFilter
    paginate_by = 10
    ordering = ["-date"]

    def get_template_names(self):
        if self.request.htmx:
            template_name = "core/registration_status_list_partial.html"
        else:
            template_name = "core/registration_status_list.html"

        return template_name

    def get(self, request, *args, **kwargs):
        if self.request.htmx:
            return super().get(request, *args, **kwargs)
        else:
            raise Http404

    def get_context_data(self, *, object_list=None, **kwargs):
        context = super().get_context_data()
        page = self.request.GET.get("page", 1)
        paginator = context["paginator"]
        context["pagination_range"] = paginator.get_elided_page_range(
            number=page, on_each_side=1, on_ends=1
        )
        return context


class ImportData(LoginRequiredMixin, PermissionRequiredMixin, TemplateView):
    template_name = "import_export.html"

    def has_permission(self):
        return self.request.user.can_import

    def post(self, request):
        def update_field(new_value, old_value):
            return new_value if not old_value else old_value

        try:
            excel_file = request.FILES["uploaded-excel"]
            sheet_obj = opx.load_workbook(excel_file).active
        except OSError:
            messages.info(request, "File is not Excel like!")
            return redirect(reverse("import-contacts"))
        except django.utils.datastructures.MultiValueDictKeyError:
            messages.info(
                request, "Please select a file then press the 'Import' button."
            )
            return redirect(reverse("import-contacts"))

        for row in sheet_obj.iter_rows(min_row=2):
            first_name = row[8].value
            last_name = row[9].value
            address = (
                row[11].value
                + "; "
                + row[12].value
                + "; "
                + row[13].value
                + "; "
                + row[14].value
                + "; "
                + row[15].value
                + "; "
                + row[16].value
            )
            focal_point = True if row[21].value == 1 else False
            org_head = True if row[23].value == 1 else False

            try:
                if first_name and last_name:
                    contact = Record.objects.get(
                        first_name=first_name, last_name=last_name
                    )
                else:
                    contact = Record.objects.get(
                        first_name=first_name, last_name=last_name, address=address
                    )
            except Record.DoesNotExist:
                Record.objects.create(
                    country=row[0].value,
                    honorific=row[2].value,
                    primary_lang=row[3].value,
                    second_lang=row[4].value,
                    third_lang=row[5].value,
                    respectful=row[6].value,
                    title=row[7].value,
                    first_name=first_name,
                    last_name=last_name,
                    designation=row[10].value,
                    address=address,
                    phones=[row[17].value],
                    faxes=[row[18].value],
                    emails=[row[19].value],
                    focal_point=focal_point,
                    org_head=org_head,
                )
            else:
                contact.country = update_field(row[0].value, contact.country)
                contact.honorific = update_field(row[2].value, contact.honorific)
                contact.primary_lang = update_field(row[3].value, contact.primary_lang)
                contact.second_lang = update_field(row[4].value, contact.second_lang)
                contact.third_lang = update_field(row[5].value, contact.third_lang)
                contact.respectful = update_field(row[6].value, contact.respectful)
                contact.title = update_field(row[7].value, contact.title)
                contact.designation = update_field(row[10].value, contact.designation)
                contact.address = update_field(address, contact.address)
                contact.phones = update_field(list(row[17].value), contact.phones)
                contact.faxes = update_field(list(row[18].value), contact.faxes)
                contact.emails = update_field(list(row[19].value), contact.emails)
                contact.focal_point = update_field(focal_point, contact.focal_point)
                contact.org_head = update_field(org_head, contact.org_head)
                contact.save()

        messages.success(request, "Successfully imported contacts.")
        return redirect(reverse("import-contacts"))


class ExportExcel(LoginRequiredMixin, View):
    def get(self, request):
        row = 1
        wb = opx.Workbook()
        ws = wb.active
        ws.title = "Contacts"
        ws.cell(row=row, column=1, value="Country name")
        ws.cell(row=row, column=2, value="Honorific")
        ws.cell(row=row, column=3, value="Respectful")
        ws.cell(row=row, column=4, value="Title")
        ws.cell(row=row, column=5, value="First Name")
        ws.cell(row=row, column=6, value="Last Name")
        ws.cell(row=row, column=7, value="Primary language")
        ws.cell(row=row, column=8, value="Second language")
        ws.cell(row=row, column=9, value="Third language")
        ws.cell(row=row, column=10, value="Organization")
        ws.cell(row=row, column=11, value="Designation")
        ws.cell(row=row, column=12, value="Head of Org.")
        ws.cell(row=row, column=13, value="Address")
        ws.cell(row=row, column=14, value="City")
        ws.cell(row=row, column=15, value="State")
        ws.cell(row=row, column=16, value="Country")
        ws.cell(row=row, column=17, value="Postal Code")
        ws.cell(row=row, column=18, value="Telephones")
        ws.cell(row=row, column=19, value="Faxes")
        ws.cell(row=row, column=20, value="E-mails")
        ws.cell(row=row, column=21, value="Focal Point")
        ws.cell(row=row, column=22, value="In Mailing List")
        ws.cell(row=row, column=23, value="Is Using Org. Address")

        for cell in ws["1:1"]:
            cell.fill = opx.styles.PatternFill(start_color="C0C0C0", fill_type="solid")
            cell.font = opx.styles.Font(bold=True, size=10, name="Arial")
            cell.border = opx.styles.Border(
                left=opx.styles.Side(border_style="thin", color="FF000000"),
                right=opx.styles.Side(border_style="thin", color="FF000000"),
            )
            cell.alignment = opx.styles.Alignment(horizontal="center")

        qs = Record.objects.prefetch_related("organization").all().order_by("pk")
        filtered = RecordFilter(request.GET, queryset=qs).qs

        for record in filtered:
            row += 1
            ws.cell(row=row, column=1, value=record.country)
            ws.cell(row=row, column=2, value=record.honorific)
            ws.cell(row=row, column=3, value=record.respectful)
            ws.cell(row=row, column=4, value=record.title)
            ws.cell(row=row, column=5, value=record.first_name)
            ws.cell(row=row, column=6, value=record.last_name)
            ws.cell(row=row, column=7, value=record.primary_lang)
            ws.cell(row=row, column=8, value=record.second_lang)
            ws.cell(row=row, column=9, value=record.third_lang)
            ws.cell(row=row, column=10, value=str(record.organization))
            ws.cell(row=row, column=11, value=record.designation)
            ws.cell(row=row, column=12, value=record.org_head)
            ws.cell(row=row, column=13, value=record.address)
            ws.cell(row=row, column=14, value=record.city)
            ws.cell(row=row, column=15, value=record.state)
            ws.cell(row=row, column=16, value=record.country)
            ws.cell(row=row, column=17, value=record.postal_code)
            ws.cell(row=row, column=18, value=str(record.phones))
            ws.cell(row=row, column=19, value=str(record.faxes))
            ws.cell(row=row, column=20, value=str(record.emails))
            ws.cell(row=row, column=21, value=record.focal_point)
            ws.cell(row=row, column=22, value=record.is_in_mailing_list)
            ws.cell(row=row, column=23, value=record.is_use_organization_address)

        for row in ws.iter_rows(min_row=2):
            for cell in row:
                cell.font = opx.styles.Font(size=10, name="Arial")
                if not cell.value:
                    cell.value = " "
        excel_file = io.BytesIO()
        wb.save(excel_file)
        response = HttpResponse(
            excel_file.getvalue(), content_type="application/ms-excel"
        )
        response["Content-Disposition"] = (
            f"attachment; filename*=UTF-8''{'Contacts.xlsx'}"
        )
        return response


class ExportDoc(LoginRequiredMixin, FilterView):
    queryset = Record.objects.prefetch_related("organization").all().order_by("pk")
    filterset_class = RecordFilter

    def get(self, request):
        doc = Document()
        section_header = doc.sections[0].header.paragraphs[0]
        section_header.text = "Ozone Contact DB\t"
        section_header.style.font.size = Pt(9)
        section_header.style.font.color.rgb = RGBColor(212, 212, 212)

        styles = doc.styles
        p_style = styles.add_style("FieldStyle", WD_STYLE_TYPE.PARAGRAPH)
        p_style.font.size = Pt(10)
        p_style.font.name = "Arial"
        p_style.font.italic = True
        p_style.paragraph_format.space_after = Pt(1.2)

        r_style = styles.add_style("RunStyle", WD_STYLE_TYPE.CHARACTER)
        r_style.font.size = Pt(10)
        r_style.font.name = "Arial"
        r_style.font.bold = True
        r_style.font.italic = False

        qs = Record.objects.prefetch_related("organization").all().order_by("pk")
        filtered = RecordFilter(request.GET, queryset=qs).qs
        ct = 1
        for record in filtered:
            doc.add_heading(str(ct) + ". " + record.first_name + " " + record.last_name)
            first_name = doc.add_paragraph("First name: ", style="FieldStyle")
            first_name.add_run(record.first_name, style="RunStyle")
            last_name = doc.add_paragraph("Last name: ", style="FieldStyle")
            last_name.add_run(record.last_name, style="RunStyle")
            title = doc.add_paragraph("Title: ", style="FieldStyle")
            title.add_run(record.title, style="RunStyle")
            honorific = doc.add_paragraph("Honorific: ", style="FieldStyle")
            honorific.add_run(record.honorific, style="RunStyle")
            respectful = doc.add_paragraph("Respectful: ", style="FieldStyle")
            respectful.add_run(record.respectful, style="RunStyle")
            primary_lang = doc.add_paragraph("Primary language: ", style="FieldStyle")
            primary_lang.add_run(record.primary_lang, style="RunStyle")
            second_lang = doc.add_paragraph("Second language: ", style="FieldStyle")
            second_lang.add_run(record.second_lang, style="RunStyle")
            third_lang = doc.add_paragraph("Third language: ", style="FieldStyle")
            third_lang.add_run(record.third_lang, style="RunStyle")
            org = doc.add_paragraph("Organization: ", style="FieldStyle")
            org.add_run(str(record.organization), style="RunStyle")
            org_head = doc.add_paragraph("Head of organization: ", style="FieldStyle")
            org_head.add_run(str(record.org_head), style="RunStyle")
            affiliation = doc.add_paragraph("Affiliation: ", style="FieldStyle")
            affiliation.add_run(record.affiliation, style="RunStyle")
            designation = doc.add_paragraph("Designation: ", style="FieldStyle")
            designation.add_run(record.designation, style="RunStyle")
            department = doc.add_paragraph("Department: ", style="FieldStyle")
            department.add_run(record.department, style="RunStyle")
            phones = doc.add_paragraph("Telephones: ", style="FieldStyle")
            phones.add_run(str(record.phones), style="RunStyle")
            faxes = doc.add_paragraph("Faxes: ", style="FieldStyle")
            faxes.add_run(str(record.faxes), style="RunStyle")
            emails = doc.add_paragraph("E-mails: ", style="FieldStyle")
            emails.add_run(str(record.emails), style="RunStyle")
            notes = doc.add_paragraph("Notes: ", style="FieldStyle")
            notes.add_run(record.notes, style="RunStyle")
            focal_point = doc.add_paragraph("Focal point: ", style="FieldStyle")
            focal_point.add_run(str(record.focal_point), style="RunStyle")
            mailing_list = doc.add_paragraph("Is in mailing list: ", style="FieldStyle")
            mailing_list.add_run(str(record.is_in_mailing_list), style="RunStyle")
            org_addr = doc.add_paragraph(
                "Is using organization address: ", style="FieldStyle"
            )
            org_addr.add_run(str(record.is_use_organization_address), style="RunStyle")
            addr = doc.add_paragraph("Address: ", style="FieldStyle")
            addr.add_run(record.address, style="RunStyle")
            city = doc.add_paragraph("City: ", style="FieldStyle")
            city.add_run(record.city, style="RunStyle")
            state = doc.add_paragraph("State: ", style="FieldStyle")
            state.add_run(record.state, style="RunStyle")
            country = doc.add_paragraph("Country: ", style="FieldStyle")
            country.add_run(record.country, style="RunStyle")
            postal_code = doc.add_paragraph("Postal code: ", style="FieldStyle")
            postal_code.add_run(record.postal_code, style="RunStyle")
            birth_date = doc.add_paragraph("Birth date: ", style="FieldStyle")
            birth_date.add_run(str(record.birth_date), style="RunStyle")
            doc.add_page_break()
            ct += 1

        doc_file = io.BytesIO()
        doc.save(doc_file)
        response = HttpResponse(doc_file.getvalue(), content_type="application/ms-word")
        response["Content-Disposition"] = (
            f"attachment; filename*=UTF-8''{'Contacts.docx'}"
        )
        return response


class GroupListView(LoginRequiredMixin, SingleTableMixin, FilterView):
    table_class = GroupTable
    queryset = Group.objects.all()
    filterset_class = GroupFilter
    paginate_by = 30

    def get_template_names(self):
        if self.request.htmx:
            template_name = "core/group_list_partial.html"
        else:
            template_name = "core/group_list.html"

        return template_name


class GroupDetailView(LoginRequiredMixin, DetailView):
    model = Group


class GroupCreateView(LoginRequiredMixin, PermissionRequiredMixin, CreateView):
    model = Group
    fields = ["name", "description"]
    template_name = "core/group_create_form.html"

    def has_permission(self):
        return self.request.user.can_edit

    def get_success_url(self):
        return reverse("group-detail", kwargs={"pk": self.object.pk})


class GroupMembersView(LoginRequiredMixin, SingleTableMixin, FilterView):
    table_class = GroupMemberTable
    queryset = Record.objects.all()
    filterset_class = GroupMembersFilter
    paginate_by = 20

    def get_template_names(self):
        if self.request.htmx:
            template_name = "core/group_members_partial.html"
        else:
            template_name = "404.html"

        return template_name

    def get(self, request, *args, **kwargs):
        if self.request.htmx:
            return super().get(request, *args, **kwargs)
        else:
            raise Http404


class GroupDeleteView(LoginRequiredMixin, PermissionRequiredMixin, DeleteView):
    model = Group

    def get_success_url(self):
        return reverse("group-list")

    def has_permission(self):
        return self.request.user.can_edit


class GroupUpdateView(LoginRequiredMixin, PermissionRequiredMixin, UpdateView):
    model = Group
    form_class = GroupUpdateForm

    def has_permission(self):
        return self.request.user.can_edit

    def get_success_url(self):
        return reverse("group-detail", kwargs={"pk": self.object.pk})


class SearchContactView(LoginRequiredMixin, FilterView, ListView):
    model = Record
    filterset_class = SearchContactFilter
    context_object_name = "contacts"
    ordering = ["first_name", "last_name"]

    def get_template_names(self):
        if self.request.htmx:
            template_name = "core/select_member_list.html"
        else:
            template_name = ""

        return template_name

    def get(self, request, *args, **kwargs):
        if self.request.htmx:
            return super().get(request, *args, **kwargs)
        else:
            raise Http404


class AddGroupMemberView(
    LoginRequiredMixin, PermissionRequiredMixin, FormMixin, DetailView
):
    template_name = "core/add_group_member.html"
    form_class = AddGroupMemberForm
    model = Group

    def has_permission(self):
        return self.request.user.can_edit

    def get_success_url(self):
        return reverse("group-detail", kwargs={"pk": self.object.pk})

    def post(self, request, *args, **kwargs):
        form = self.get_form()
        self.object = self.get_object()

        if form.is_valid():
            self.new_member = Record.objects.filter(
                pk=form.cleaned_data["new_member_id"]
            ).first()
            if self.new_member is not None:
                return self.form_valid(form)

        return self.form_invalid(form)

    def form_valid(self, form):
        self.object.contacts.add(self.new_member)
        self.object.save()
        return super().form_valid(form)


class AddMultipleGroupMembersView(
    LoginRequiredMixin, PermissionRequiredMixin, FilterView, FormMixin
):
    form_class = AddMultipleGroupMembersForm
    model = Record
    template_name = "core/add_multiple_group_members.html"
    filterset_class = RecordFilter

    def has_permission(self):
        return self.request.user.can_edit

    def get_success_url(self):
        return reverse("group-list")

    def get_template_names(self):
        if self.request.htmx:
            template_name = "core/add_multiple_group_members_partial.html"
        else:
            template_name = "core/add_multiple_group_members.html"

        return template_name

    def post(self, request, *args, **kwargs):
        self.object_list = self.get_filterset(self.filterset_class).qs
        form = self.get_form()

        if form.is_valid():
            return self.form_valid(form)

        return self.form_invalid(form)

    def form_valid(self, form):
        members = Record.objects.filter(id__in=form.cleaned_data["members"])
        groups = Group.objects.filter(id__in=form.cleaned_data["groups"])
        for group in groups:
            group.contacts.add(*members.all())
        return super().form_valid(form)

    def form_invalid(self, form):
        f = self.get_filterset(self.filterset_class)
        previously_selected_members = []
        member_ids = form.cleaned_data.get("members")
        if member_ids and len(member_ids) > 0:
            previously_selected_members = Record.objects.filter(id__in=member_ids)
        return self.render_to_response(
            self.get_context_data(
                form=form,
                filter=f,
                previously_selected_members=previously_selected_members,
            )
        )


class EmailTemplateCreateView(LoginRequiredMixin, PermissionRequiredMixin, CreateView):
    model = EmailTemplate
    fields = "__all__"

    def has_permission(self):
        return self.request.user.can_send_mail

    def get_success_url(self):
        return reverse("email-template-create-success")

    def get_template_names(self):
        if self.request.htmx:
            return "core/create_email_template.html"
        return "404.html"

    def get(self, request, *args, **kwargs):
        if self.request.htmx:
            return super().get(request, *args, **kwargs)
        raise Http404


class EmailTemplateCreateSuccessView(
    LoginRequiredMixin, PermissionRequiredMixin, TemplateView
):
    def has_permission(self):
        return self.request.user.can_send_mail

    def get_template_names(self):
        if self.request.htmx:
            return "core/create_email_template_success.html"
        return "404.html"

    def get(self, request, *args, **kwargs):
        if self.request.htmx:
            return super().get(request, *args, **kwargs)
        raise Http404


class EmailPage(LoginRequiredMixin, PermissionRequiredMixin, FormMixin, FilterView):
    template_name = "send_emails.html"
    form_class = SendEmailForm
    filterset_class = RecordFilter

    def has_permission(self):
        return self.request.user.can_send_mail

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)

        context["email_templates"] = EmailTemplate.objects.all()

        return context

    def log_sent_email(
        self,
        subject,
        content,
        recipients,
        cc,
        send_personalised_emails=False,
        groups=None,
        cc_groups=None,
    ):
        email = Emails.objects.create(
            subject=subject,
            content=content,
            send_personalised_emails=send_personalised_emails,
        )
        if groups:
            email.groups.add(*groups)

        if cc_groups:
            email.groups.add(*cc_groups)

        email.recipients.set(recipients)
        if cc:
            email.cc.set(cc)

        return email

    def post(self, request, true=None, *args, **kwargs):
        form = self.get_form()
        form.is_valid()

        contacts = []
        groups_ids = form.cleaned_data.get("groups")
        members = form.cleaned_data.get("members")
        groups = None

        if groups_ids:
            groups = Group.objects.filter(id__in=groups_ids)
            contacts = Record.objects.filter(group__in=groups)
        elif members:
            contacts = Record.objects.filter(id__in=members)

        cc_contacts = []
        cc_groups_ids = form.cleaned_data.get("cc_groups")
        cc_ids = form.cleaned_data.get("cc")

        cc_groups = None

        if cc_groups_ids:
            cc_groups = Group.objects.filter(id__in=cc_groups_ids)
            cc_contacts = Record.objects.filter(group__in=cc_groups)

        elif cc_ids:
            cc_contacts = Record.objects.filter(id__in=cc_ids)

        subject = form.cleaned_data.get("subject")
        content = form.cleaned_data.get("content")
        files = form.cleaned_data.get("files")

        if not contacts:
            messages.info(request, "No contact selected!")
            return redirect(reverse("emails-page"))
        if not subject:
            messages.info(request, "No subject selected!")
            return redirect(reverse("emails-page"))
        if not content:
            messages.info(request, "No content selected!")
            return redirect(reverse("emails-page"))

        if form.cleaned_data["send_personalised_emails"]:
            email = self.log_sent_email(
                form.cleaned_data["subject"],
                form.cleaned_data["content"],
                contacts,
                None,
                form.cleaned_data["send_personalised_emails"],
                groups,
                None,
            )
        else:
            email = self.log_sent_email(
                form.cleaned_data["subject"],
                form.cleaned_data["content"],
                contacts,
                cc_contacts,
                form.cleaned_data["send_personalised_emails"],
                groups,
                cc_groups,
            )
        if files:
            for f in files:
                print(f)
                EmailFile.objects.create(
                    name=f.name,
                    file=f,
                    email=email,
                )
        SendMailTask.objects.create(email=email).run(is_async=true)

        messages.success(request, "Successfully created task for sending emails!")
        return redirect(reverse("emails-page"))


class SyncKronosView(LoginRequiredMixin, PermissionRequiredMixin, TemplateView):
    template_name = "core/sync_kronos.html"

    def has_permission(self):
        return self.request.user.can_import

    def get_context_data(self, **kwargs):
        context = super().get_context_data()
        context["load_kronos_events_tasks"] = LoadKronosEventsTask.objects.order_by(
            "-created_on"
        )
        return context


class RunKronosEventsImport(LoginRequiredMixin, PermissionRequiredMixin, FormView):
    form_class = KronosEventsImportForm

    def get_success_url(self):
        return reverse("kronos-events-import")

    def get_context_data(self, **kwargs):
        context = super().get_context_data()
        running_tasks = LoadKronosEventsTask.objects.filter(
            status__in=LoadKronosEventsTask.TASK_STATUS_PENDING_VALUES
        ).exists()
        if running_tasks:
            context["kronos_events_importing"] = True
        return context

    def has_permission(self):
        return self.request.user.can_import

    def get_template_names(self):
        if self.request.htmx:
            template_name = "core/import_kronos_events.html"
        else:
            template_name = ""

        return template_name

    def get(self, request, *args, **kwargs):
        if self.request.htmx:
            return super().get(request, *args, **kwargs)
        else:
            raise Http404

    def form_valid(self, form):
        LoadKronosEventsTask.objects.create(created_by=self.request.user).run(
            is_async=True
        )
        return super().form_valid(form)


class LoadKronosEventsView(LoginRequiredMixin, SingleTableView):
    table_class = LoadKronosEventsTable
    queryset = LoadKronosEventsTask.objects.all().order_by("-started_on")
    paginate_by = 10

    def get_template_names(self):
        if self.request.htmx:
            template_name = "core/kronos_events_tasks_partial.html"
        else:
            template_name = "404.html"

        return template_name

    def get(self, request, *args, **kwargs):
        if self.request.htmx:
            return super().get(request, *args, **kwargs)
        else:
            raise Http404


class RunKronosParticipantsImport(
    LoginRequiredMixin, PermissionRequiredMixin, FormView
):
    form_class = KronosParticipantsImportForm

    def get_success_url(self):
        return reverse("kronos-participants-import")

    def get_context_data(self, **kwargs):
        context = super().get_context_data()
        running_tasks = LoadKronosParticipantsTask.objects.filter(
            status__in=LoadKronosParticipantsTask.TASK_STATUS_PENDING_VALUES
        ).exists()
        if running_tasks:
            context["kronos_participants_importing"] = True

        if TemporaryContact.objects.exists():
            context["conflicts"] = True

        context["events"] = KronosEvent.objects.all()

        return context

    def has_permission(self):
        return self.request.user.can_import

    def get_template_names(self):
        if self.request.htmx:
            template_name = "core/import_kronos_participants.html"
        else:
            template_name = "404.html"

        return template_name

    def get(self, request, *args, **kwargs):
        if self.request.htmx:
            return super().get(request, *args, **kwargs)
        else:
            raise Http404

    def form_valid(self, form):
        event_ids = form.cleaned_data.get("events")
        events = KronosEvent.objects.filter(id__in=event_ids)
        task = LoadKronosParticipantsTask.objects.create(
            created_by=self.request.user,
            create_groups=form.cleaned_data.get("create_groups"),
        )
        task.kronos_events.set(events)
        task.run(is_async=True)
        return super().form_valid(form)


class LoadKronosParticipantsView(LoginRequiredMixin, SingleTableView):
    table_class = LoadKronosParticipantsTable
    queryset = LoadKronosParticipantsTask.objects.all().order_by("-started_on")
    paginate_by = 10

    def get_template_names(self):
        if self.request.htmx:
            template_name = "core/kronos_participants_tasks_partial.html"
        else:
            template_name = "404.html"

        return template_name

    def get(self, request, *args, **kwargs):
        if self.request.htmx:
            return super().get(request, *args, **kwargs)
        else:
            raise Http404


class ResolveConflictsView(LoginRequiredMixin, PermissionRequiredMixin, ListView):
    model = TemporaryContact
    queryset = TemporaryContact.objects.select_related(
        "record", "organization", "record__organization"
    )
    paginate_by = 30
    ordering = ["first_name", "last_name"]

    def get_context_data(self, *, object_list=None, **kwargs):
        context = super().get_context_data()
        page = self.request.GET.get("page", 1)
        paginator = context["paginator"]
        context["pagination_range"] = paginator.get_elided_page_range(
            number=page, on_each_side=1, on_ends=1
        )
        return context

    def has_permission(self):
        return self.request.user.can_import

    def get_template_names(self):
        if self.request.htmx:
            return "core/resolve_conflicts_partial.html"
        return "core/resolve_conflicts.html"

    def get(self, request, *args, **kwargs):
        running_tasks = ResolveAllConflictsTask.objects.filter(
            status__in=LoadKronosEventsTask.TASK_STATUS_PENDING_VALUES
        ).exists()
        if running_tasks:
            return redirect(reverse("conflicts-resolving"))
        if not TemporaryContact.objects.exists() and not self.request.htmx:
            return redirect(reverse("no-conflicts"))
        return super().get(request, *args, **kwargs)


class ResolveConflictsFormView(LoginRequiredMixin, PermissionRequiredMixin, FormView):
    form_class = ResolveConflictForm

    def has_permission(self):
        return self.request.user.can_import

    def get_context_data(self, **kwargs):
        context = super().get_context_data()

        if self.request.GET.get("conflict"):
            contact = TemporaryContact.objects.filter(
                id=self.request.GET.get("conflict")
            ).first()
            if contact:
                context["contact"] = contact

        if (
            self.request.GET.get("method")
            == ConflictResolutionMethods.KEEP_OLD_DATA.value
        ):
            context["KEEP_OLD_DATA"] = True
        elif (
            self.request.GET.get("method")
            == ConflictResolutionMethods.SAVE_INCOMING_DATA.value
        ):
            context["SAVE_INCOMING_DATA"] = True
        return context

    def get_success_url(self):
        return reverse("conflict-resolved")

    def get_template_names(self):
        if self.request.htmx:
            template_name = "core/resolve_conflict.html"
        else:
            template_name = ""

        return template_name

    def get_initial(self):
        initial = super().get_initial()
        if self.request.GET.get("conflict"):
            initial["incoming_contact"] = self.request.GET.get("conflict")

        if (
            self.request.GET.get("method")
            == ConflictResolutionMethods.KEEP_OLD_DATA.value
        ):
            initial["method"] = ConflictResolutionMethods.KEEP_OLD_DATA.value

        if (
            self.request.GET.get("method")
            == ConflictResolutionMethods.SAVE_INCOMING_DATA.value
        ):
            initial["method"] = ConflictResolutionMethods.SAVE_INCOMING_DATA.value

        return initial

    def get(self, request, *args, **kwargs):
        if self.request.htmx:
            return super().get(request, *args, **kwargs)
        else:
            raise Http404

    def form_valid(self, form):
        if form.cleaned_data.get("method") == ConflictResolutionMethods.KEEP_OLD_DATA:
            incoming_contact_id = form.cleaned_data.get("incoming_contact")
            TemporaryContact.objects.filter(id=incoming_contact_id).delete()
        elif (
            form.cleaned_data.get("method")
            == ConflictResolutionMethods.SAVE_INCOMING_DATA
        ):
            incoming_contact_id = form.cleaned_data.get("incoming_contact")
            incoming_contact = (
                TemporaryContact.objects.filter(id=incoming_contact_id)
                .select_related("record")
                .first()
            )
            record = incoming_contact.record
            update_values = vars(incoming_contact)
            incoming_contact.delete()
            update_values.pop("record_id")
            update_values.pop("id")
            update_object(record, update_values)

        return super().form_valid(form)


class ConflictsResolvedView(LoginRequiredMixin, PermissionRequiredMixin, TemplateView):
    template_name = "core/conflict_resolved.html"

    def has_permission(self):
        return self.request.user.can_import


class ResolveAllConflictsFormView(
    LoginRequiredMixin, PermissionRequiredMixin, FormView
):
    form_class = ResolveAllConflictsForm

    def has_permission(self):
        return self.request.user.can_import

    def get_context_data(self, **kwargs):
        context = super().get_context_data()
        if (
            self.request.GET.get("method")
            == ConflictResolutionMethods.KEEP_OLD_DATA.value
        ):
            context["KEEP_OLD_DATA"] = True
        elif (
            self.request.GET.get("method")
            == ConflictResolutionMethods.SAVE_INCOMING_DATA.value
        ):
            context["SAVE_INCOMING_DATA"] = True
        return context

    def get_success_url(self):
        return reverse("all-conflicts-resolved")

    def get_template_names(self):
        if self.request.htmx:
            template_name = "core/resolve_all_conflicts.html"
        else:
            template_name = ""

        return template_name

    def get_initial(self):
        initial = super().get_initial()

        if (
            self.request.GET.get("method")
            == ConflictResolutionMethods.KEEP_OLD_DATA.value
        ):
            initial["method"] = ConflictResolutionMethods.KEEP_OLD_DATA.value
        elif (
            self.request.GET.get("method")
            == ConflictResolutionMethods.SAVE_INCOMING_DATA.value
        ):
            initial["method"] = ConflictResolutionMethods.SAVE_INCOMING_DATA.value

        return initial

    def get(self, request, *args, **kwargs):
        if self.request.htmx:
            return super().get(request, *args, **kwargs)
        else:
            raise Http404

    def form_valid(self, form):
        ResolveAllConflictsTask.objects.create(
            created_by=self.request.user, method=form.cleaned_data.get("method")
        ).run(is_async=True)
        return super().form_valid(form)


class AllConflictsResolvedView(
    LoginRequiredMixin, PermissionRequiredMixin, TemplateView
):
    template_name = "core/all_conflicts_resolved.html"

    def has_permission(self):
        return self.request.user.can_import


class ConflictsResolvingView(LoginRequiredMixin, PermissionRequiredMixin, TemplateView):
    def get_context_data(self, **kwargs):
        context = super().get_context_data()

        running_tasks = ResolveAllConflictsTask.objects.filter(
            status__in=LoadKronosEventsTask.TASK_STATUS_PENDING_VALUES
        )
        if len(running_tasks) > 0:
            context["conflicts_resolving"] = True
        return context

    def has_permission(self):
        return self.request.user.can_import

    def get_template_names(self):
        if self.request.htmx:
            template_name = "core/conflicts_resolving_partial.html"
        else:
            template_name = "core/conflicts_resolving.html"

        return template_name


class NoConflictsView(LoginRequiredMixin, PermissionRequiredMixin, TemplateView):
    template_name = "core/no_conflicts.html"

    def has_permission(self):
        return self.request.user.can_import


class MergeContactsView(LoginRequiredMixin, PermissionRequiredMixin, TemplateView):
    template_name = "core/merge_conflicts.html"

    def has_permission(self):
        return self.request.user.can_edit


class MergeContactsFirstStepView(LoginRequiredMixin, PermissionRequiredMixin, FormView):
    form_class = MergeContactsFirstStepForm

    def has_permission(self):
        return self.request.user.can_edit

    def get_template_names(self):
        if self.request.htmx:
            template_name = "core/merge_contacts_first_step.html"
        else:
            template_name = "404.html"

        return template_name

    def get(self, request, *args, **kwargs):
        if self.request.htmx:
            return super().get(request, *args, **kwargs)
        else:
            raise Http404

    def form_valid(self, form):
        ids = form.cleaned_data["contacts"]
        return HttpResponseRedirect(
            reverse("merge-second-step")
            + f"?ids={','.join(str(contact_id) for contact_id in ids)}"
        )


class MergeContactsSecondStepView(
    LoginRequiredMixin, PermissionRequiredMixin, FormView
):
    form_class = MergeContactsSecondStepForm

    def get_success_url(self):
        return reverse("merge-success")

    def has_permission(self):
        return self.request.user.can_edit

    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        kwargs["choices"] = self.get_form_contact_choices()
        return kwargs

    def get_form_contact_choices(self):
        ids = self.request.GET.get("ids")

        if ids:
            ids = ids.split(",")
            choices = [
                (
                    contact.get("id"),
                    contact.get("first_name") + " " + contact.get("last_name"),
                )
                for contact in Record.objects.filter(id__in=ids).values(
                    "id", "first_name", "last_name"
                )
            ]
            return choices
        return []

    def get_context_data(self, **kwargs):
        context = super().get_context_data()
        form = context.get("form")
        if form:
            contact_ids = [choice[0] for choice in form.fields["contact"].choices]
            context["contacts"] = Record.objects.filter(id__in=contact_ids)
        return context

    def get_template_names(self):
        if self.request.htmx:
            template_name = "core/merge_contacts_second_step.html"
        else:
            template_name = "404.html"

        return template_name

    def get(self, request, *args, **kwargs):
        if self.request.htmx:
            return super().get(request, *args, **kwargs)
        else:
            raise Http404

    def form_valid(self, form):
        selected_contact = Record.objects.filter(
            id=form.cleaned_data["contact"]
        ).first()
        secondary_ids = []
        if self.request.GET.get("ids"):
            secondary_ids = [
                contact_id
                for contact_id in self.request.GET.get("ids").split(",")
                if int(contact_id) != int(form.cleaned_data["contact"])
            ]
        secondary_contacts = Record.objects.filter(id__in=secondary_ids)
        for secondary_contact in secondary_contacts:
            secondary_contact.main_contact = selected_contact
            secondary_contact.save()
            if secondary_contact.record_set.exists():
                for contact in secondary_contact.record_set.all():
                    contact.main_contact = selected_contact
                    contact.save()

        return super().form_valid(form)


class MergeSuccessView(LoginRequiredMixin, PermissionRequiredMixin, TemplateView):
    def has_permission(self):
        return self.request.user.can_edit

    def get_template_names(self):
        if self.request.htmx:
            template_name = "core/merge_success.html"
        else:
            template_name = "404.html"

        return template_name

    def get(self, request, *args, **kwargs):
        if self.request.htmx:
            return super().get(request, *args, **kwargs)
        else:
            raise Http404


class EmailListView(LoginRequiredMixin, FilterView, ListView):
    model = Emails
    paginate_by = 30
    filterset_class = EmailFilter
    ordering = ["-created_at"]

    def get_context_data(self, *, object_list=None, **kwargs):
        context = super().get_context_data(object_list=None, **kwargs)
        page = self.request.GET.get("page", 1)
        paginator = context["paginator"]
        context["pagination_range"] = paginator.get_elided_page_range(
            number=page, on_each_side=1, on_ends=1
        )
        return context

    def get_template_names(self):
        if self.request.htmx:
            template_name = "core/email_list.html"
        else:
            template_name = "core/emails_history.html"

        return template_name


class ContactEmailsHistory(LoginRequiredMixin, DetailView):
    model = Record
    template_name = "core/contact_emails_history.html"


class GroupEmailsHistory(LoginRequiredMixin, DetailView):
    model = Group
    template_name = "core/group_emails_history.html"


class EmailDetailView(LoginRequiredMixin, DetailView):
    model = Emails
    template_name = "core/email_detail.html"
