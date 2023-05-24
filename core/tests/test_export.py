import io
import pytest
import openpyxl

from docx import Document
from django.urls import reverse

pytestmark = [pytest.mark.django_db]


def test_export_docx(login_user_can_edit, contact, other_contact):
    client, user = login_user_can_edit
    contact.organization.save()
    contact.first_name = "First"
    contact.last_name = "Last"
    contact.save()
    other_contact.organization.save()
    other_contact.save()

    url = reverse("export-docx")
    response = client.get(url)
    assert response.status_code == 200

    doc = Document(io.BytesIO(response.content))
    assert doc.paragraphs[0].text == "1. First Last"
    assert doc.paragraphs[28].text == "2. Other Record"


def test_export_docx_filtered(login_user_can_edit, contact, other_contact):
    client, user = login_user_can_edit
    contact.organization.save()
    contact.save()
    other_contact.organization.save()
    other_contact.save()

    url = reverse("export-docx")
    response = client.get(
        url,
        {
            "name": "Other",
        },
    )
    assert response.status_code == 200

    doc = Document(io.BytesIO(response.content))
    assert doc.paragraphs[0].text == "1. Other Record"


def test_export_excel(login_user_can_edit, contact, other_contact):
    client, user = login_user_can_edit
    contact.organization.save()
    contact.first_name = "First"
    contact.last_name = "Last"
    contact.save()
    other_contact.organization.save()
    other_contact.save()

    url = reverse("export-excel")
    response = client.get(url)
    assert response.status_code == 200

    excel = openpyxl.load_workbook(io.BytesIO(response.content))
    sheet = excel.active
    assert sheet.cell(row=2, column=5).value == "First"
    assert sheet.cell(row=2, column=6).value == "Last"
    assert sheet.cell(row=3, column=5).value == "Other"
    assert sheet.cell(row=3, column=6).value == "Record"


def test_export_excel_filtered(login_user_can_edit, contact, other_contact):
    client, user = login_user_can_edit
    contact.organization.save()
    contact.save()
    other_contact.organization.save()
    other_contact.save()

    url = reverse("export-excel")
    response = client.get(
        url,
        {
            "name": "Other",
        },
    )
    assert response.status_code == 200
    excel = openpyxl.load_workbook(io.BytesIO(response.content))
    sheet = excel.active
    assert sheet.cell(row=2, column=5).value == "Other"
    assert sheet.cell(row=2, column=6).value == "Record"
